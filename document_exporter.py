"""
文档导出工具模块
支持将分析结果导出为Markdown和DOCX格式
"""
from typing import List, Dict, Any, Optional
from datetime import datetime
from io import BytesIO
import re

from logger import get_logger
from config import settings

logger = get_logger(__name__)


class DocumentExporter:
    """文档导出器类"""
    
    @staticmethod
    def export_to_markdown(
        event_text: str,
        analysis: str,
        retrieved_cards: List[Dict[str, Any]],
        references: Optional[List[str]] = None,
        llm_model: str = "DeepSeek"
    ) -> str:
        """
        导出为Markdown格式
        
        Args:
            event_text: 事件文本
            analysis: 分析结果
            retrieved_cards: 检索到的理论卡片
            references: 引用的理论ID列表
            llm_model: 使用的LLM模型
        
        Returns:
            Markdown格式文本
        """
        md_content = f"""# 不曾社科理论RAG分析系统 - 分析报告

> 本分析报告由北域工作室开发的BucengRAG分析系统生成，基于 {llm_model} 模型生成。
>
> 开发者: 北域工作室 (Northland Comprehensive Studio)
> 官网: https://beiyu.xuanjian.top/

---

生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}

---

## 一、事件描述

{event_text}

---

## 二、理论分析

{analysis}

---

## 三、引用的理论卡片

"""
        
        if retrieved_cards:
            for i, card in enumerate(retrieved_cards, 1):
                md_content += f"""### {i}. {card['title']}

**ID**: {card['id']}

**相似度**: {card.get('similarity', 0):.2%}

**内容**:
{card['content']}

**来源**: {card['source']}

**关键词**: {', '.join(card.get('keywords', []))}

---

"""
        
        md_content += f"""
---

*本报告由不曾社科理论RAG分析系统生成*

*开发者: 北域工作室 (Northland Comprehensive Studio)*

*官网: https://beiyu.xuanjian.top/*

*模型: {llm_model}*

*生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}*
"""
        
        logger.info("Markdown文档生成成功")
        return md_content
    
    @staticmethod
    def export_to_docx(
        event_text: str,
        analysis: str,
        retrieved_cards: List[Dict[str, Any]],
        references: Optional[List[str]] = None,
        llm_model: str = "DeepSeek"
    ) -> BytesIO:
        """
        导出为DOCX格式
        
        Args:
            event_text: 事件文本
            analysis: 分析结果
            retrieved_cards: 检索到的理论卡片
            references: 引用的理论ID列表
            llm_model: 使用的LLM模型
        
        Returns:
            DOCX文件的BytesIO对象
        """
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            doc = Document()
            
            title = doc.add_heading('不曾社科理论RAG分析系统 - 分析报告', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            info_para = doc.add_paragraph()
            info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            info_run = info_para.add_run(f'本分析报告由北域工作室开发的BucengRAG分析系统生成，基于 {llm_model} 模型生成。')
            info_run.font.size = Pt(10)
            info_run.font.color.rgb = RGBColor(100, 100, 100)
            info_run.italic = True
            
            dev_para = doc.add_paragraph()
            dev_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            dev_run = dev_para.add_run('开发者: 北域工作室 (Northland Comprehensive Studio) | 官网: https://beiyu.xuanjian.top/')
            dev_run.font.size = Pt(9)
            dev_run.font.color.rgb = RGBColor(128, 128, 128)
            
            doc.add_paragraph()
            
            time_para = doc.add_paragraph()
            time_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            time_run = time_para.add_run(f'生成时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
            time_run.font.size = Pt(10)
            time_run.font.color.rgb = RGBColor(128, 128, 128)
            
            doc.add_paragraph()
            doc.add_paragraph('=' * 50)
            
            doc.add_heading('一、事件描述', 1)
            doc.add_paragraph(event_text)
            
            doc.add_paragraph()
            doc.add_paragraph('=' * 50)
            
            doc.add_heading('二、理论分析', 1)
            
            analysis_lines = analysis.split('\n')
            for line in analysis_lines:
                line = line.strip()
                if line:
                    if line.startswith('## '):
                        doc.add_heading(line[3:], 2)
                    elif line.startswith('### '):
                        doc.add_heading(line[4:], 3)
                    elif line.startswith('#### '):
                        doc.add_heading(line[5:], 4)
                    else:
                        para = doc.add_paragraph()
                        parts = re.split(r'(\*\*[^*]+\*\*)', line)
                        for part in parts:
                            if part.startswith('**') and part.endswith('**'):
                                run = para.add_run(part[2:-2])
                                run.bold = True
                            else:
                                para.add_run(part)
            
            doc.add_paragraph()
            doc.add_paragraph('=' * 50)
            
            doc.add_heading('三、引用的理论卡片', 1)
            
            if retrieved_cards:
                for i, card in enumerate(retrieved_cards, 1):
                    doc.add_heading(f'{i}. {card["title"]}', 2)
                    
                    id_para = doc.add_paragraph()
                    id_run = id_para.add_run(f'ID: {card["id"]}')
                    id_run.font.size = Pt(10)
                    id_run.font.color.rgb = RGBColor(100, 100, 100)
                    
                    similarity_para = doc.add_paragraph()
                    similarity_run = similarity_para.add_run(f'相似度: {card.get("similarity", 0):.2%}')
                    similarity_run.font.size = Pt(10)
                    similarity_run.font.color.rgb = RGBColor(100, 100, 100)
                    
                    doc.add_paragraph()
                    content_para = doc.add_paragraph()
                    content_run = content_para.add_run('内容: ')
                    content_run.bold = True
                    content_para.add_run(card['content'])
                    
                    source_para = doc.add_paragraph()
                    source_run = source_para.add_run('来源: ')
                    source_run.bold = True
                    source_para.add_run(card['source'])
                    
                    if card.get('keywords'):
                        keywords_para = doc.add_paragraph()
                        keywords_run = keywords_para.add_run('关键词: ')
                        keywords_run.bold = True
                        keywords_para.add_run(', '.join(card['keywords']))
                    
                    doc.add_paragraph()
            
            doc.add_paragraph('=' * 50)
            
            footer_para = doc.add_paragraph()
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            footer_run = footer_para.add_run('本报告由不曾社科理论RAG分析系统生成')
            footer_run.font.size = Pt(9)
            footer_run.font.color.rgb = RGBColor(128, 128, 128)
            footer_run.italic = True
            
            dev_footer = doc.add_paragraph()
            dev_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
            dev_run = dev_footer.add_run('开发者: 北域工作室 (Northland Comprehensive Studio)')
            dev_run.font.size = Pt(9)
            dev_run.font.color.rgb = RGBColor(128, 128, 128)
            
            url_footer = doc.add_paragraph()
            url_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
            url_run = url_footer.add_run('官网: https://beiyu.xuanjian.top/')
            url_run.font.size = Pt(9)
            url_run.font.color.rgb = RGBColor(128, 128, 128)
            
            model_para = doc.add_paragraph()
            model_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            model_run = model_para.add_run(f'模型: {llm_model} | 生成时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
            model_run.font.size = Pt(9)
            model_run.font.color.rgb = RGBColor(128, 128, 128)
            
            docx_buffer = BytesIO()
            doc.save(docx_buffer)
            docx_buffer.seek(0)
            
            logger.info("DOCX文档生成成功")
            return docx_buffer
            
        except Exception as e:
            logger.error(f"DOCX文档生成失败: {e}")
            raise
    
    @staticmethod
    def get_export_filename(format_type: str, prefix: str = "analysis") -> str:
        """
        生成导出文件名
        
        Args:
            format_type: 格式类型 (md/docx)
            prefix: 文件名前缀
        
        Returns:
            文件名
        """
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        extension = 'md' if format_type == 'md' else 'docx'
        return f"{prefix}_{timestamp}.{extension}"


def export_analysis_result(
    event_text: str,
    analysis: str,
    retrieved_cards: List[Dict[str, Any]],
    format_type: str = 'md',
    llm_model: str = None
) -> tuple:
    """
    导出分析结果
    
    Args:
        event_text: 事件文本
        analysis: 分析结果
        retrieved_cards: 检索到的理论卡片
        format_type: 格式类型 (md/docx)
        llm_model: 使用的LLM模型
    
    Returns:
        (文件内容, 文件名)
    """
    exporter = DocumentExporter()
    
    if llm_model is None:
        llm_model = settings.LLM_MODEL
    
    if format_type == 'md':
        content = exporter.export_to_markdown(event_text, analysis, retrieved_cards, llm_model=llm_model)
        filename = exporter.get_export_filename('md')
        return content, filename
    
    elif format_type == 'docx':
        content = exporter.export_to_docx(event_text, analysis, retrieved_cards, llm_model=llm_model)
        filename = exporter.get_export_filename('docx')
        return content, filename
    
    else:
        raise ValueError(f"不支持的格式类型: {format_type}")
