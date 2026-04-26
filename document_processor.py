"""
文档处理工具模块
支持PDF、DOCX、TXT、MD等格式的文档解析
"""
from typing import List, Dict, Any, Optional
from pathlib import Path
import re

from logger import get_logger

logger = get_logger(__name__)


class DocumentProcessor:
    """文档处理器类"""
    
    @staticmethod
    def extract_text_from_pdf(file_path: str) -> str:
        """
        从PDF文件提取文本
        
        Args:
            file_path: PDF文件路径
        
        Returns:
            提取的文本内容
        """
        try:
            import pdfplumber
            
            text = ""
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            
            logger.info(f"成功从PDF提取文本: {file_path}, 共{len(text)}字符")
            return text
            
        except Exception as e:
            logger.error(f"PDF文本提取失败: {e}")
            raise
    
    @staticmethod
    def extract_text_from_docx(file_path: str) -> str:
        """
        从DOCX文件提取文本
        
        Args:
            file_path: DOCX文件路径
        
        Returns:
            提取的文本内容
        """
        try:
            from docx import Document
            
            doc = Document(file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # 提取表格内容
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + "\t"
                    text += "\n"
            
            logger.info(f"成功从DOCX提取文本: {file_path}, 共{len(text)}字符")
            return text
            
        except Exception as e:
            logger.error(f"DOCX文本提取失败: {e}")
            raise
    
    @staticmethod
    def extract_text_from_txt(file_path: str) -> str:
        """
        从TXT文件提取文本
        
        Args:
            file_path: TXT文件路径
        
        Returns:
            提取的文本内容
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            
            logger.info(f"成功从TXT提取文本: {file_path}, 共{len(text)}字符")
            return text
            
        except Exception as e:
            logger.error(f"TXT文本提取失败: {e}")
            raise
    
    @staticmethod
    def extract_text_from_md(file_path: str) -> str:
        """
        从Markdown文件提取文本
        
        Args:
            file_path: Markdown文件路径
        
        Returns:
            提取的文本内容
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            
            # 简单的Markdown清理
            # 移除代码块
            text = re.sub(r'```[\s\S]*?```', '', text)
            # 移除行内代码
            text = re.sub(r'`[^`]+`', '', text)
            # 移除链接但保留文本
            text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
            # 移除图片
            text = re.sub(r'!\[([^\]]*)\]\([^\)]+\)', '', text)
            # 移除标题标记
            text = re.sub(r'^#{1,6}\s+', '', text, flags=re.MULTILINE)
            # 移除粗体和斜体标记
            text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
            text = re.sub(r'\*([^*]+)\*', r'\1', text)
            
            logger.info(f"成功从Markdown提取文本: {file_path}, 共{len(text)}字符")
            return text
            
        except Exception as e:
            logger.error(f"Markdown文本提取失败: {e}")
            raise
    
    @classmethod
    def extract_text(cls, file_path: str) -> str:
        """
        根据文件类型自动提取文本
        
        Args:
            file_path: 文件路径
        
        Returns:
            提取的文本内容
        
        Raises:
            ValueError: 不支持的文件格式
        """
        path = Path(file_path)
        suffix = path.suffix.lower()
        
        if suffix == '.pdf':
            return cls.extract_text_from_pdf(file_path)
        elif suffix in ['.docx', '.doc']:
            return cls.extract_text_from_docx(file_path)
        elif suffix == '.txt':
            return cls.extract_text_from_txt(file_path)
        elif suffix == '.md':
            return cls.extract_text_from_md(file_path)
        else:
            raise ValueError(f"不支持的文件格式: {suffix}")
    
    @staticmethod
    def split_text_into_chunks(
        text: str,
        chunk_size: int = 500,
        overlap: int = 50
    ) -> List[str]:
        """
        将文本分割成块
        
        Args:
            text: 原始文本
            chunk_size: 每块大小（字符数）
            overlap: 重叠大小
        
        Returns:
            文本块列表
        """
        # 按段落分割
        paragraphs = text.split('\n\n')
        
        chunks = []
        current_chunk = ""
        
        for para in paragraphs:
            para = para.strip()
            if not para:
                continue
            
            # 如果当前块加上新段落不超过限制，则添加
            if len(current_chunk) + len(para) + 2 <= chunk_size:
                current_chunk += para + "\n\n"
            else:
                # 保存当前块
                if current_chunk:
                    chunks.append(current_chunk.strip())
                
                # 如果段落本身超过限制，需要进一步分割
                if len(para) > chunk_size:
                    # 按句子分割
                    sentences = re.split(r'[。！？\n]', para)
                    temp_chunk = ""
                    
                    for sentence in sentences:
                        sentence = sentence.strip()
                        if not sentence:
                            continue
                        
                        if len(temp_chunk) + len(sentence) + 1 <= chunk_size:
                            temp_chunk += sentence + "。"
                        else:
                            if temp_chunk:
                                chunks.append(temp_chunk.strip())
                            temp_chunk = sentence + "。"
                    
                    if temp_chunk:
                        current_chunk = temp_chunk + "\n\n"
                    else:
                        current_chunk = ""
                else:
                    current_chunk = para + "\n\n"
        
        # 保存最后一块
        if current_chunk:
            chunks.append(current_chunk.strip())
        
        logger.info(f"文本分割完成，共{len(chunks)}块")
        return chunks


def extract_text_from_uploaded_file(uploaded_file) -> str:
    """
    从Streamlit上传的文件提取文本
    
    Args:
        uploaded_file: Streamlit UploadedFile对象
    
    Returns:
        提取的文本内容
    """
    import tempfile
    import os
    
    # 创建临时文件
    with tempfile.NamedTemporaryFile(delete=False, suffix=Path(uploaded_file.name).suffix) as tmp:
        tmp.write(uploaded_file.getvalue())
        tmp_path = tmp.name
    
    try:
        # 提取文本
        text = DocumentProcessor.extract_text(tmp_path)
        return text
    finally:
        # 删除临时文件
        os.unlink(tmp_path)


def process_document_for_knowledge_base(
    text: str,
    source_name: str,
    chunk_size: int = 500
) -> List[Dict[str, Any]]:
    """
    处理文档并准备导入知识库的卡片
    
    Args:
        text: 文档文本
        source_name: 来源名称
        chunk_size: 文本块大小
    
    Returns:
        卡片列表
    """
    # 分割文本
    chunks = DocumentProcessor.split_text_into_chunks(text, chunk_size)
    
    # 创建卡片
    cards = []
    for i, chunk in enumerate(chunks):
        card = {
            'id': f"{source_name}_chunk_{i}",
            'title': f"{source_name} - 片段 {i+1}",
            'content': chunk,
            'keywords': [],
            'source': source_name
        }
        cards.append(card)
    
    return cards
